#MESSAGE : EVERY POINT IS EXPLAINED WITH THE HELP OF COMMENTS
#MOREOVER : IMPORTING FILES IS AN ANOTHER TOPIC
import os
import numpy as np
import pandas as pd
import docx

#creating a folder which will store our final result
directory = "Master_folder" # folder name
p_dir = "C:/Users/hp/OneDrive/Desktop/" 
# this is the path to desktop on my computer - you can have different one
path = os.path.join(p_dir, directory)
# if folder is already created
try: 
    os.mkdir(path) 
except OSError as e: 
    print("Folder Already Exists")

#print("Directory " ,directory, "Created")
# Now All the work we have to do is whatever our output will be we have to send it to the created folder(which we created)

#getting the data into our program
df = pd.read_excel("data.xlsx")
name = df.columns[1];
company = df[name][0]
address = df[name][1]
number = df[name][2]
lis = df[name][3]
#setting the data into temp files
#filling the highlighted part of the temp files
#opening temp files
temp_1 = docx.Document('temp_file_1.docx') # opening first temp file
temp_2 = docx.Document('temp_file_2.docx') # opening the 2nd temp file
temp_3 = docx.Document('temp_file_3.docx')
temp_4 = docx.Document('temp_file_4.docx')
temp_5 = docx.Document('temp_file_5.docx')


#WRITING TO THEM AT HIGHLIGHTED PART
temp_1.paragraphs[0].text = name
temp_1.paragraphs[1].text = company
temp_1.paragraphs[2].text = address
temp_1.paragraphs[3].text = number
temp_1.paragraphs[4].text = lis
#first temp file done
# you have used a table so to access the table elements(text...etc) see below code
temp_2.tables[0].cell(0,1).text =  name
temp_2.tables[0].cell(1,1).text = company
temp_2.tables[0].cell(2,1).text = address
temp_2.tables[0].cell(3,1).text = number
temp_2.tables[0].cell(4,1).text = lis
#2nd file done
for p in temp_3.paragraphs:
    p.text = p.text.replace("company",company)
    p.text = p.text.replace("address",address)
    p.text = p.text.replace("number",number)
# third file done

for p in temp_4.paragraphs:
    p.text = p.text.replace("company",company)
# fourth file done

temp_5.paragraphs[0].text = address
temp_5.paragraphs[1].text = company
temp_5.paragraphs[2].text = lis
temp_5.paragraphs[3].text = name
temp_5.paragraphs[4].text = number
#fifth file done


#saving the changes - you can use path as on your computer it may vary computer to computer
temp_1.save('C:/Users/hp/OneDrive/Desktop/Master_folder/temp_1.docx')
temp_2.save('C:/Users/hp/OneDrive/Desktop/Master_folder/temp_2.docx')
temp_3.save('C:/Users/hp/OneDrive/Desktop/Master_folder/temp_3.docx')

temp_4.save('C:/Users/hp/OneDrive/Desktop/Master_folder/temp_4.docx')
temp_5.save('C:/Users/hp/OneDrive/Desktop/Master_folder/temp_5.docx')
#print(df['Ankush Jain'][])
#company = df['company'][1]





